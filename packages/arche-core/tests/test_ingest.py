"""Tests for file ingestion — text extraction from multiple formats."""

from __future__ import annotations

import json
from pathlib import Path

import pytest

from arche.workflow._ingest import extract_text


# ═══════════════════════════════════════════════════════════════════════════════
# TXT extraction
# ═══════════════════════════════════════════════════════════════════════════════


def test_extract_txt(tmp_path: Path) -> None:
    """Plain text files are read as-is with normalized whitespace."""
    f = tmp_path / "sample.txt"
    f.write_text("Fatima Abdullahi, NIN 12345678901\n", encoding="utf-8")
    result = extract_text(f)
    assert "Fatima Abdullahi" in result
    assert "12345678901" in result


def test_extract_txt_unicode(tmp_path: Path) -> None:
    """Unicode text (African names, accented chars) is preserved."""
    f = tmp_path / "unicode.txt"
    f.write_text("Ngozi Okonkwo-Eze, Abuja\nFatoumata Diallo, Dakar\n", encoding="utf-8")
    result = extract_text(f)
    assert "Ngozi Okonkwo-Eze" in result
    assert "Fatoumata Diallo" in result


def test_extract_txt_whitespace_normalization(tmp_path: Path) -> None:
    """Excessive whitespace is collapsed; paragraph breaks preserved."""
    f = tmp_path / "messy.txt"
    f.write_text("Hello\n\n\n\n\nWorld\n\n  Trailing spaces  \n", encoding="utf-8")
    result = extract_text(f)
    # 3+ newlines collapsed to 2
    assert "\n\n\n" not in result
    assert "Hello\n\nWorld" in result


def test_extract_empty_txt(tmp_path: Path) -> None:
    """Empty files return an empty string."""
    f = tmp_path / "empty.txt"
    f.write_text("", encoding="utf-8")
    result = extract_text(f)
    assert result == ""


def test_extract_whitespace_only_txt(tmp_path: Path) -> None:
    """Files with only whitespace return empty string."""
    f = tmp_path / "whitespace.txt"
    f.write_text("   \n\n   \t  \n", encoding="utf-8")
    result = extract_text(f)
    assert result == ""


# ═══════════════════════════════════════════════════════════════════════════════
# CSV extraction
# ═══════════════════════════════════════════════════════════════════════════════


def test_extract_csv(tmp_path: Path) -> None:
    """CSV files are read as plain text."""
    f = tmp_path / "records.csv"
    f.write_text("name,nin,phone\nFatima,12345678901,+234 803 555 7890\n", encoding="utf-8")
    result = extract_text(f)
    assert "Fatima" in result
    assert "12345678901" in result


# ═══════════════════════════════════════════════════════════════════════════════
# JSON extraction
# ═══════════════════════════════════════════════════════════════════════════════


def test_extract_json(tmp_path: Path) -> None:
    """JSON files are read as plain text."""
    f = tmp_path / "data.json"
    data = {"name": "Amina Bello", "phone": "+234 803 111 2222"}
    f.write_text(json.dumps(data, indent=2), encoding="utf-8")
    result = extract_text(f)
    assert "Amina Bello" in result
    assert "+234 803 111 2222" in result


# ═══════════════════════════════════════════════════════════════════════════════
# Markdown / log files
# ═══════════════════════════════════════════════════════════════════════════════


def test_extract_md(tmp_path: Path) -> None:
    """Markdown files are supported as plain text."""
    f = tmp_path / "notes.md"
    f.write_text("# Patient Notes\n\nName: Chukwuma Obi\n", encoding="utf-8")
    result = extract_text(f)
    assert "Chukwuma Obi" in result


def test_extract_log(tmp_path: Path) -> None:
    """Log files are supported as plain text."""
    f = tmp_path / "app.log"
    f.write_text("[INFO] Patient registered: Yemi Alade\n", encoding="utf-8")
    result = extract_text(f)
    assert "Yemi Alade" in result


# ═══════════════════════════════════════════════════════════════════════════════
# PDF extraction (optional dependency)
# ═══════════════════════════════════════════════════════════════════════════════


_has_pymupdf = False
try:
    import fitz  # noqa: F401
    _has_pymupdf = True
except ImportError:
    pass


@pytest.mark.skipif(not _has_pymupdf, reason="pymupdf not installed")
def test_extract_pdf(tmp_path: Path) -> None:
    """PDF files are extracted via pymupdf."""
    import fitz

    f = tmp_path / "test.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Fatima Abdullahi NIN 12345678901")
    doc.save(str(f))
    doc.close()

    result = extract_text(f)
    assert "Fatima" in result
    assert "12345678901" in result


@pytest.mark.skipif(not _has_pymupdf, reason="pymupdf not installed")
def test_extract_pdf_multipage(tmp_path: Path) -> None:
    """Multi-page PDFs join pages with newlines."""
    import fitz

    f = tmp_path / "multi.pdf"
    doc = fitz.open()
    for i, text in enumerate(["Page one content", "Page two content"]):
        page = doc.new_page()
        page.insert_text((72, 72), text)
    doc.save(str(f))
    doc.close()

    result = extract_text(f)
    assert "Page one content" in result
    assert "Page two content" in result


# ═══════════════════════════════════════════════════════════════════════════════
# DOCX extraction (optional dependency)
# ═══════════════════════════════════════════════════════════════════════════════


_has_docx = False
try:
    from docx import Document  # noqa: F401
    _has_docx = True
except ImportError:
    pass


@pytest.mark.skipif(not _has_docx, reason="python-docx not installed")
def test_extract_docx(tmp_path: Path) -> None:
    """DOCX files are extracted via python-docx."""
    from docx import Document

    f = tmp_path / "test.docx"
    doc = Document()
    doc.add_paragraph("Fatima Abdullahi NIN 12345678901")
    doc.add_paragraph("Phone: +234 803 555 7890")
    doc.save(str(f))

    result = extract_text(f)
    assert "Fatima Abdullahi" in result
    assert "+234 803 555 7890" in result


@pytest.mark.skipif(not _has_docx, reason="python-docx not installed")
def test_extract_docx_empty_paragraphs(tmp_path: Path) -> None:
    """Empty paragraphs in DOCX are skipped."""
    from docx import Document

    f = tmp_path / "sparse.docx"
    doc = Document()
    doc.add_paragraph("First")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("Second")
    doc.save(str(f))

    result = extract_text(f)
    assert "First" in result
    assert "Second" in result


# ═══════════════════════════════════════════════════════════════════════════════
# Error handling
# ═══════════════════════════════════════════════════════════════════════════════


def test_unsupported_format_raises_valueerror(tmp_path: Path) -> None:
    """Unsupported file extensions raise ValueError with helpful message."""
    f = tmp_path / "data.xyz"
    f.write_text("some data", encoding="utf-8")
    with pytest.raises(ValueError, match="Unsupported file format"):
        extract_text(f)


def test_file_not_found_raises() -> None:
    """Non-existent files raise FileNotFoundError."""
    with pytest.raises(FileNotFoundError, match="File not found"):
        extract_text("/nonexistent/file.txt")


def test_string_path_accepted(tmp_path: Path) -> None:
    """String paths are accepted in addition to Path objects."""
    f = tmp_path / "string_path.txt"
    f.write_text("Test content", encoding="utf-8")
    result = extract_text(str(f))
    assert result == "Test content"


# ═══════════════════════════════════════════════════════════════════════════════
# Public API
# ═══════════════════════════════════════════════════════════════════════════════


def test_extract_text_in_public_api() -> None:
    """``extract_text`` is importable from the top-level ``arche`` package.

    The top-level name lazily resolves to its real home
    (``arche.workflow._ingest``), so this back-compat surface stays available
    without emitting a DeprecationWarning.
    """
    from arche import extract_text as et
    assert callable(et)
